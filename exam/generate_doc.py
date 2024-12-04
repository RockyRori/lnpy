from docx import Document


# 生成考点文档
def create_checkpoints_doc():
    doc = Document()
    doc.add_heading('考点总结 - CDS 534: Database Management Lecture 1', 0)
    doc.add_paragraph('1. 课程概述：数据管理的高级概念，包括数据库类型、数据库管理系统、数据建模、查询语言等。')
    doc.add_paragraph('2. 数据库类型：关系型数据库与非关系型数据库（SQL与NoSQL）。')
    doc.add_paragraph('3. 数据管理与数据安全：数据灾难恢复、隐私控制、用户安全等基本概念与最佳实践。')
    doc.add_paragraph('4. 数据的管理与分析：如何通过数据管理系统进行数据处理、分析、可视化与报告生成。')
    doc.add_paragraph('5. SQL与NoSQL：理解如何使用SQL与NoSQL进行数据库操作，分析不同数据库的应用场景。')
    doc.add_paragraph('6. 数据库设计与开发：如何进行数据库设计、数据库建模以及开发数据库管理系统来解决实际问题。')

    return doc


# 生成问题文档
def create_questions_doc():
    doc = Document()
    doc.add_heading('问题 - CDS 534: Database Management Lecture 1', 0)

    # 单项选择题
    doc.add_paragraph('单项选择题: 以下哪项不属于CDS 534课程的学习目标？')
    doc.add_paragraph(
        'A. 使用SQL和NoSQL查询数据库\nB. 理解数据灾难恢复与规划\nC. 进行高级算法建模与编程\nD. 开发数据库管理系统来解决现实问题')
    doc.add_paragraph('答案: C')

    doc.add_paragraph('单项选择题: 数据库设计过程中，哪种数据模型通常使用表格来表示数据？')
    doc.add_paragraph('A. 非结构化数据模型\nB. 关系型数据模型\nC. 图形数据模型\nD. 文档数据模型')
    doc.add_paragraph('答案: B')

    # 填空题
    doc.add_paragraph('填空题: 关系型数据库通常使用__________来描述数据表之间的关系。')
    doc.add_paragraph('答案: 外键')

    doc.add_paragraph('填空题: 在数据灾难恢复中，______是指在发生系统故障后恢复数据的过程。')
    doc.add_paragraph('答案: 备份与恢复')

    # 判断题
    doc.add_paragraph('判断题: SQL是一种常用的非关系型数据库查询语言。')
    doc.add_paragraph('答案: False')

    doc.add_paragraph('判断题: 数据管理系统不仅可以用于数据的存储和检索，还可以用于数据的可视化与分析。')
    doc.add_paragraph('答案: True')

    # 计算题
    doc.add_paragraph(
        '计算题: 假设有两个表格，`Employee` 和 `Department`，请写出SQL查询，返回所有员工的姓名和他们所属的部门名称。')
    doc.add_paragraph('答案:')
    doc.add_paragraph("""
    SELECT E.Name, D.DepartmentName
    FROM Employee E
    JOIN Department D ON E.DepartmentID = D.DepartmentID;
    """)

    # 简答题
    doc.add_paragraph('简答题: 请简述SQL与NoSQL数据库的主要区别及应用场景。')
    doc.add_paragraph(
        '答案: SQL数据库是结构化的，适用于事务型操作和复杂查询（如关系型数据库）。NoSQL数据库是非结构化的，适用于大数据和快速扩展的应用，如文档存储、键值对存储等。')

    # 编程题
    doc.add_paragraph('编程题: 使用Python编写一个程序，连接MySQL数据库并查询员工表中的所有数据。')
    doc.add_paragraph('答案:')
    doc.add_paragraph("""
    import mysql.connector

    # 连接数据库
    conn = mysql.connector.connect(
        host="localhost",
        user="yourusername",
        password="yourpassword",
        database="yourdatabase"
    )

    cursor = conn.cursor()

    # 执行查询
    cursor.execute("SELECT * FROM Employee")

    # 获取结果
    for row in cursor.fetchall():
        print(row)

    # 关闭连接
    conn.close()
    """)

    # SQL题
    doc.add_paragraph(
        'SQL题: 使用SQL创建一个`Employee`表，包含以下字段：员工ID（整数主键）、员工姓名（文本）、部门ID（整数外键）、入职日期（日期）。')
    doc.add_paragraph('答案:')
    doc.add_paragraph("""
    CREATE TABLE Employee (
        EmployeeID INT PRIMARY KEY,
        Name TEXT,
        DepartmentID INT,
        HireDate DATE,
        FOREIGN KEY (DepartmentID) REFERENCES Department(DepartmentID)
    );
    """)

    return doc


# 生成答案文档
def create_answers_doc():
    doc = Document()
    doc.add_heading('答案 - CDS 534: Database Management Lecture 1', 0)

    doc.add_paragraph('单项选择题: 以下哪项不属于CDS 534课程的学习目标？')
    doc.add_paragraph('答案: C')

    doc.add_paragraph('单项选择题: 数据库设计过程中，哪种数据模型通常使用表格来表示数据？')
    doc.add_paragraph('答案: B')

    doc.add_paragraph('填空题: 关系型数据库通常使用__________来描述数据表之间的关系。')
    doc.add_paragraph('答案: 外键')

    doc.add_paragraph('填空题: 在数据灾难恢复中，______是指在发生系统故障后恢复数据的过程。')
    doc.add_paragraph('答案: 备份与恢复')

    doc.add_paragraph('判断题: SQL是一种常用的非关系型数据库查询语言。')
    doc.add_paragraph('答案: False')

    doc.add_paragraph('判断题: 数据管理系统不仅可以用于数据的存储和检索，还可以用于数据的可视化与分析。')
    doc.add_paragraph('答案: True')

    doc.add_paragraph(
        '计算题: 假设有两个表格，`Employee` 和 `Department`，请写出SQL查询，返回所有员工的姓名和他们所属的部门名称。')
    doc.add_paragraph('答案:')
    doc.add_paragraph("""
    SELECT E.Name, D.DepartmentName
    FROM Employee E
    JOIN Department D ON E.DepartmentID = D.DepartmentID;
    """)

    doc.add_paragraph('简答题: 请简述SQL与NoSQL数据库的主要区别及应用场景。')
    doc.add_paragraph(
        '答案: SQL数据库是结构化的，适用于事务型操作和复杂查询（如关系型数据库）。NoSQL数据库是非结构化的，适用于大数据和快速扩展的应用，如文档存储、键值对存储等。')

    doc.add_paragraph('编程题: 使用Python编写一个程序，连接MySQL数据库并查询员工表中的所有数据。')
    doc.add_paragraph('答案:')
    doc.add_paragraph("""
    import mysql.connector

    # 连接数据库
    conn = mysql.connector.connect(
        host="localhost",
        user="yourusername",
        password="yourpassword",
        database="yourdatabase"
    )

    cursor = conn.cursor()

    # 执行查询
    cursor.execute("SELECT * FROM Employee")

    # 获取结果
    for row in cursor.fetchall():
        print(row)

    # 关闭连接
    conn.close()
    """)

    doc.add_paragraph(
        'SQL题: 使用SQL创建一个`Employee`表，包含以下字段：员工ID（整数主键）、员工姓名（文本）、部门ID（整数外键）、入职日期（日期）。')
    doc.add_paragraph('答案:')
    doc.add_paragraph("""
    CREATE TABLE Employee (
        EmployeeID INT PRIMARY KEY,
        Name TEXT,
        DepartmentID INT,
        HireDate DATE,
        FOREIGN KEY (DepartmentID) REFERENCES Department(DepartmentID)
    );
    """)

    return doc


# 保存文档
checkpoints_doc = create_checkpoints_doc()
questions_doc = create_questions_doc()
answers_doc = create_answers_doc()

checkpoints_doc.save("考点.docx")
questions_doc.save("问题.docx")
answers_doc.save("答案.docx")

"考点.docx, 问题.docx, 答案.docx"
